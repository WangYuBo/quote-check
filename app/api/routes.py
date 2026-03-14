"""
FastAPI route definitions for the proofreader web application.

Endpoints
---------
GET  /                       – Serve the main HTML page (Jinja2 template)
POST /api/proofread          – Accept file uploads, enqueue background task,
                               return task_id
GET  /api/result/{task_id}  – Poll task status / retrieve completed result
GET  /api/stream/{task_id}  – SSE stream of task progress
GET  /api/export/{task_id}  – Export result as docx or csv
GET  /health                – Simple liveness probe
"""

import asyncio
import csv
import io
import logging
import time
import uuid
from pathlib import Path
from typing import AsyncGenerator

from fastapi import APIRouter, BackgroundTasks, File, HTTPException, Request, UploadFile
from fastapi.responses import HTMLResponse, Response, StreamingResponse
from fastapi.templating import Jinja2Templates

from app.config import settings
from app.services.file_parser import (
    FileNotFoundError as ParserFileNotFoundError,
    FileParseError,
    UnsupportedFileTypeError,
    parse_file,
    parse_manuscript_structured,
)
from app.services.proofreader import run_proofreading

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Router & templates
# ---------------------------------------------------------------------------

router = APIRouter()

# Templates instance is configured here; the search path is resolved relative
# to the project root at import time.
_TEMPLATES_DIR = Path(__file__).parent.parent.parent / "templates"
templates = Jinja2Templates(directory=str(_TEMPLATES_DIR))

# ---------------------------------------------------------------------------
# In-memory task store
# task_id (str) -> {
#     "status":           "pending" | "processing" | "done" | "error",
#     "result":           dict | None,
#     "error":            str  | None,
#     "logs":             list[str],
#     "source_filenames": list[str],
#     "created_at":       float,
# }
# ---------------------------------------------------------------------------

_tasks: dict[str, dict] = {}


# ---------------------------------------------------------------------------
# Background worker
# ---------------------------------------------------------------------------

async def _run_task(
    task_id: str,
    manuscript_path: str,
    manuscript_ext: str,
    source_paths: list[tuple[str, str, str]],
) -> None:
    """Parse all files and run the proofreading pipeline.

    Parameters
    ----------
    task_id:
        Unique identifier for this task.
    manuscript_path:
        Filesystem path to the saved manuscript file.
    manuscript_ext:
        Lowercase file extension of the manuscript (e.g. ``.docx``).
    source_paths:
        List of ``(path, ext, original_filename)`` tuples for each source file.

    Updates _tasks[task_id] with the final status.
    """
    _tasks[task_id]["status"] = "processing"
    logger.info("[task %s] 开始处理", task_id)

    def log_cb(message: str) -> None:
        timestamp = time.strftime("%H:%M:%S", time.localtime())
        entry = f"[{timestamp}] {message}"
        _tasks[task_id]["logs"].append(entry)
        logger.debug("[task %s] %s", task_id, message)

    try:
        paragraphs = parse_manuscript_structured(manuscript_path, manuscript_ext)

        source_texts: dict[str, str] = {}
        for path, ext, original_filename in source_paths:
            text = parse_file(path, ext)
            source_texts[original_filename] = text

        result = await run_proofreading(paragraphs, source_texts, log_callback=log_cb)

        _tasks[task_id]["status"] = "done"
        _tasks[task_id]["result"] = result
        logger.info("[task %s] 校对完成", task_id)

    except (ParserFileNotFoundError, UnsupportedFileTypeError, FileParseError) as exc:
        logger.error("[task %s] 文件解析失败: %s", task_id, exc)
        _tasks[task_id]["status"] = "error"
        _tasks[task_id]["error"] = str(exc)

    except Exception as exc:  # pylint: disable=broad-except
        logger.error("[task %s] 未知错误: %s", task_id, exc, exc_info=True)
        _tasks[task_id]["status"] = "error"
        _tasks[task_id]["error"] = f"校对过程发生未知错误：{exc}"


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@router.get("/", response_class=HTMLResponse, summary="主页")
async def index(request: Request) -> HTMLResponse:
    """Render and return the main application page."""
    return templates.TemplateResponse("index.html", {"request": request})


@router.get("/health", summary="健康检查")
async def health() -> dict:
    """Return a simple alive response for load-balancer / monitoring checks."""
    return {"status": "ok"}


@router.post("/api/proofread", summary="提交校对任务")
async def proofread(
    background_tasks: BackgroundTasks,
    manuscript: UploadFile = File(..., description="书稿文件（docx / pdf / md / txt）"),
    sources: list[UploadFile] = File(..., description="参考原文（1-10个文件，docx / pdf / md / txt）"),
) -> dict:
    """Accept uploaded files and start an async proofreading task.

    Returns
    -------
    JSON
        ``{"task_id": "<uuid>", "status": "pending"}``
    """
    if not sources or len(sources) < 1:
        raise HTTPException(status_code=400, detail="至少需要上传一个参考原文文件。")
    if len(sources) > 10:
        raise HTTPException(status_code=400, detail="参考原文文件数量不能超过10个。")

    settings.ensure_upload_dir()

    task_id = str(uuid.uuid4())

    # --- Validate & persist manuscript ---
    manuscript_suffix = _validated_suffix(manuscript.filename, "manuscript")
    manuscript_path = settings.upload_dir / f"{task_id}_manuscript{manuscript_suffix}"
    await _save_upload(manuscript, manuscript_path)

    # --- Validate & persist source files ---
    source_paths: list[tuple[str, str, str]] = []
    source_filenames: list[str] = []

    for i, source in enumerate(sources):
        source_suffix = _validated_suffix(source.filename, f"source[{i}]")
        source_path = settings.upload_dir / f"{task_id}_source_{i}{source_suffix}"
        await _save_upload(source, source_path)
        original_filename = source.filename or f"source_{i}{source_suffix}"
        source_paths.append((str(source_path), source_suffix, original_filename))
        source_filenames.append(original_filename)

    # --- Register task ---
    _tasks[task_id] = {
        "status": "pending",
        "result": None,
        "error": None,
        "logs": [],
        "source_filenames": source_filenames,
        "created_at": time.time(),
    }

    # --- Enqueue background processing ---
    background_tasks.add_task(
        _run_task,
        task_id,
        str(manuscript_path),
        manuscript_suffix,
        source_paths,
    )

    logger.info(
        "[task %s] 任务已注册，书稿=%s，原文=%s",
        task_id,
        manuscript.filename,
        source_filenames,
    )

    return {"task_id": task_id, "status": "pending"}


@router.get("/api/result/{task_id}", summary="查询校对结果")
async def get_result(task_id: str) -> dict:
    """Return the current status and (when available) result of a task.

    Returns
    -------
    JSON
        | ``{"task_id": "…", "status": "pending", "logs": [...]}``
        | ``{"task_id": "…", "status": "processing", "logs": [...]}``
        | ``{"task_id": "…", "status": "done", "result": {…}, "missing_sources": [...]}``
        | ``{"task_id": "…", "status": "error", "error": "…"}``

    Raises
    ------
    404
        If *task_id* is not found.
    """
    task = _tasks.get(task_id)
    if task is None:
        raise HTTPException(status_code=404, detail=f"Task '{task_id}' not found.")

    response: dict = {"task_id": task_id, "status": task["status"]}

    if task["status"] == "done":
        response["result"] = task["result"]
        result = task["result"] or {}
        response["missing_sources"] = result.get("missing_sources", [])
    elif task["status"] == "error":
        response["error"] = task["error"]
    elif task["status"] in ("processing", "pending"):
        response["logs"] = task.get("logs", [])

    return response


@router.get("/api/stream/{task_id}", summary="SSE 任务进度流")
async def stream_task(task_id: str) -> StreamingResponse:
    """Stream task progress as Server-Sent Events.

    Pushes a JSON event every second until the task is done or errored.

    Raises
    ------
    404
        If *task_id* is not found.
    """
    if task_id not in _tasks:
        raise HTTPException(status_code=404, detail=f"Task '{task_id}' not found.")

    async def event_generator() -> AsyncGenerator[str, None]:
        import json

        while True:
            task = _tasks.get(task_id)
            if task is None:
                break

            status = task["status"]
            logs = task.get("logs", [])

            if status == "done":
                result = task.get("result") or {}
                payload = {
                    "status": "done",
                    "result": task["result"],
                    "logs": logs,
                    "missing_sources": result.get("missing_sources", []),
                }
                yield f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"
                break
            elif status == "error":
                payload = {
                    "status": "error",
                    "error": task.get("error"),
                    "logs": logs,
                }
                yield f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"
                break
            else:
                payload = {
                    "status": status,
                    "logs": logs,
                }
                yield f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"

            await asyncio.sleep(1)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@router.get("/api/export/{task_id}", summary="导出校对结果")
async def export_result(task_id: str, format: str = "docx") -> Response:
    """Export a completed proofreading result as a downloadable file.

    Parameters
    ----------
    task_id:
        The task whose result to export.
    format:
        ``"docx"`` (default) or ``"csv"``.

    Raises
    ------
    404
        If *task_id* is not found.
    400
        If the task is not yet done, or if *format* is unsupported.
    """
    task = _tasks.get(task_id)
    if task is None:
        raise HTTPException(status_code=404, detail=f"Task '{task_id}' not found.")

    if task["status"] != "done":
        raise HTTPException(
            status_code=400,
            detail=f"任务尚未完成（当前状态：{task['status']}），无法导出。",
        )

    if format not in ("docx", "csv"):
        raise HTTPException(
            status_code=400,
            detail=f"不支持的导出格式 '{format}'，请使用 'docx' 或 'csv'。",
        )

    result = task.get("result") or {}
    items: list[dict] = result.get("results", [])

    columns = ["序号", "引用文字", "校对结果", "问题描述", "综合评价"]

    def _build_rows() -> list[list[str]]:
        rows = []
        for idx, item in enumerate(items, start=1):
            quote = item.get("quote", "")
            verdict = item.get("verdict", "")
            issues = " | ".join(
                filter(
                    None,
                    [
                        item.get("text_issues", ""),
                        item.get("explanation_issues", ""),
                        item.get("context_issues", ""),
                    ],
                )
            )
            summary = item.get("summary", "")
            rows.append([str(idx), quote, verdict, issues, summary])
        return rows

    if format == "csv":
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(columns)
        writer.writerows(_build_rows())
        csv_bytes = output.getvalue().encode("utf-8-sig")  # BOM for Excel compatibility
        return Response(
            content=csv_bytes,
            media_type="text/csv",
            headers={
                "Content-Disposition": f'attachment; filename="proofreader_result_{task_id}.csv"'
            },
        )

    # format == "docx"
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("校对结果", level=1)

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(columns):
        hdr_cells[i].text = col_name
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(11)

    # Data rows
    for row_data in _build_rows():
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text

    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="proofreader_result_{task_id}.docx"'
        },
    )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _validated_suffix(filename: str | None, field_name: str) -> str:
    """Return the lowercase file extension or raise HTTP 400."""
    if not filename:
        raise HTTPException(
            status_code=400,
            detail=f"'{field_name}' 文件名为空，无法判断文件类型。",
        )
    suffix = Path(filename).suffix.lower()
    if suffix not in settings.supported_extensions:
        raise HTTPException(
            status_code=400,
            detail=(
                f"'{field_name}' 文件类型 '{suffix}' 不被支持。"
                f"支持的格式：{settings.supported_extensions}"
            ),
        )
    return suffix


async def _save_upload(upload: UploadFile, dest: Path) -> None:
    """Stream an UploadFile to *dest*, enforcing the configured size limit."""
    max_bytes = settings.max_upload_size_bytes
    written = 0

    with dest.open("wb") as fh:
        while chunk := await upload.read(1024 * 64):  # 64 KB chunks
            written += len(chunk)
            if written > max_bytes:
                dest.unlink(missing_ok=True)
                raise HTTPException(
                    status_code=413,
                    detail=(
                        f"文件超过大小限制 {settings.max_upload_size_mb} MB，"
                        "请压缩后重新上传。"
                    ),
                )
            fh.write(chunk)
