"""
File parser service.

Parses uploaded files (.docx, .pdf, .md, .txt) into plain text strings
that downstream services (proofreader, quote extractor, etc.) can consume.
"""

import os
import re
from pathlib import Path


def _clean_text(text: str) -> str:
    """清理文本中的不可见 Unicode 控制字符（零宽字符、软连字符等）。"""
    # 移除零宽字符、软连字符等不可见字符
    text = re.sub(r"[\u200b\u200c\u200d\u200e\u200f\u00ad\ufeff\u2028\u2029]", "", text)
    # 规范化空白（保留换行，压缩多余空格）
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Public exceptions
# ---------------------------------------------------------------------------

class FileNotFoundError(Exception):  # noqa: A001 – intentional shadowing
    """Raised when the requested file does not exist on disk."""


class UnsupportedFileTypeError(Exception):
    """Raised when the file extension is not supported."""


class FileParseError(Exception):
    """Raised when a supported file cannot be parsed (e.g. corrupt file)."""


# ---------------------------------------------------------------------------
# Supported extension → parser mapping
# ---------------------------------------------------------------------------

_SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".md", ".txt", ".epub"}


# ---------------------------------------------------------------------------
# Internal parsers
# ---------------------------------------------------------------------------

def _parse_docx(file_path: str) -> str:
    """Extract all paragraph text from a .docx file."""
    try:
        from docx import Document  # python-docx
    except ImportError as exc:
        raise FileParseError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from exc

    try:
        doc = Document(file_path)
    except Exception as exc:
        raise FileParseError(f"Failed to open docx file '{file_path}': {exc}") from exc

    paragraphs = [_clean_text(para.text) for para in doc.paragraphs if para.text.strip()]
    return "\n".join(p for p in paragraphs if p)


def parse_docx_structured(file_path: str) -> dict:
    """Parse a .docx file and return structured paragraph data with chapter info.

    Returns a dict with:
    - ``paragraphs``: list of dicts with index, text, style, is_heading, chapter
    - ``plain_text``: full plain text (for backward compat)
    """
    try:
        from docx import Document  # python-docx
    except ImportError as exc:
        raise FileParseError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from exc

    try:
        doc = Document(file_path)
    except Exception as exc:
        raise FileParseError(f"Failed to open docx file '{file_path}': {exc}") from exc

    result_paragraphs = []
    current_chapter = ""
    idx = 0
    cumulative_chars = 0
    # 每页约 900 个中文字符（A4 标准版面估算）
    chars_per_page = 900

    for para in doc.paragraphs:
        text = _clean_text(para.text)
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        is_heading = "heading" in style_name.lower() or "标题" in style_name

        if is_heading:
            current_chapter = text

        cumulative_chars += len(text)
        estimated_page = (cumulative_chars // chars_per_page) + 1

        idx += 1
        result_paragraphs.append({
            "index": idx,
            "text": text,
            "style": style_name,
            "is_heading": is_heading,
            "chapter": current_chapter,
            "estimated_page": estimated_page,
        })

    plain_text = "\n".join(p["text"] for p in result_paragraphs)
    return {"paragraphs": result_paragraphs, "plain_text": plain_text}


def _parse_pdf(file_path: str) -> str:
    """Extract text from every page of a .pdf file using pdfplumber."""
    try:
        import pdfplumber
    except ImportError as exc:
        raise FileParseError(
            "pdfplumber is not installed. Run: pip install pdfplumber"
        ) from exc

    pages_text: list[str] = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
    except FileParseError:
        raise
    except Exception as exc:
        raise FileParseError(f"Failed to parse PDF '{file_path}': {exc}") from exc

    return "\n".join(pages_text)


def _parse_epub(file_path: str) -> str:
    """Extract text from an .epub file using ebooklib and BeautifulSoup."""
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise FileParseError("ebooklib/beautifulsoup4 not installed. Run: pip install EbookLib beautifulsoup4") from exc

    try:
        book = epub.read_epub(file_path)
    except Exception as exc:
        raise FileParseError(f"Failed to open epub file '{file_path}': {exc}") from exc

    chapters_text = []
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            text = soup.get_text('\n')
            text = _clean_text(text)
            if text:
                chapters_text.append(text)

    return "\n\n".join(chapters_text)


def _parse_text(file_path: str) -> str:
    """Read a plain-text file (.md or .txt) and return its contents."""
    try:
        with open(file_path, "r", encoding="utf-8") as fh:
            return fh.read()
    except UnicodeDecodeError:
        # Fallback: try system default encoding
        try:
            with open(file_path, "r", encoding="latin-1") as fh:
                return fh.read()
        except Exception as exc:
            raise FileParseError(
                f"Could not decode text file '{file_path}': {exc}"
            ) from exc
    except Exception as exc:
        raise FileParseError(f"Failed to read file '{file_path}': {exc}") from exc


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_file(file_path: str, file_type: str | None = None) -> str:
    """Parse a file and return its full text content.

    Parameters
    ----------
    file_path:
        Absolute or relative path to the file on disk.
    file_type:
        Optional explicit extension (e.g. ``".docx"``).  When omitted the
        extension is inferred from *file_path*.

    Returns
    -------
    str
        Plain-text representation of the file content.

    Raises
    ------
    FileNotFoundError
        If *file_path* does not point to an existing file.
    UnsupportedFileTypeError
        If the extension is not in the supported set.
    FileParseError
        If the file exists and has a supported format but cannot be parsed.
    """
    path = Path(file_path)

    # --- Guard: file must exist ---
    if not path.exists():
        raise FileNotFoundError(f"File not found: '{file_path}'")
    if not path.is_file():
        raise FileNotFoundError(f"Path is not a regular file: '{file_path}'")

    # --- Determine extension ---
    ext = (file_type or path.suffix).lower()
    if not ext.startswith("."):
        ext = f".{ext}"

    if ext not in _SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'. "
            f"Supported types: {sorted(_SUPPORTED_EXTENSIONS)}"
        )

    # --- Dispatch to the right parser ---
    if ext == ".docx":
        return _parse_docx(str(path))
    if ext == ".pdf":
        return _parse_pdf(str(path))
    if ext in {".md", ".txt"}:
        return _parse_text(str(path))
    if ext == ".epub":
        return _parse_epub(str(path))

    # Should never reach here given the guard above, but keeps mypy happy.
    raise UnsupportedFileTypeError(f"No parser registered for extension '{ext}'")


def parse_manuscript_structured(file_path: str, file_type: str | None = None) -> list[dict]:
    """Parse a manuscript file and return a list of paragraph dicts.

    For .docx files, uses structural metadata (headings → chapter names).
    For other formats, splits on double newlines with no chapter info.

    Each dict contains: index, text, chapter, is_heading.
    """
    path = Path(file_path)
    if not path.exists() or not path.is_file():
        raise FileNotFoundError(f"File not found: '{file_path}'")

    ext = (file_type or path.suffix).lower()
    if not ext.startswith("."):
        ext = f".{ext}"

    if ext == ".docx":
        structured = parse_docx_structured(str(path))
        return structured["paragraphs"]

    # For PDF/MD/TXT: get plain text then split into paragraphs
    plain_text = parse_file(file_path, file_type)
    raw_paras = [p.strip() for p in plain_text.split("\n\n") if p.strip()]
    if not raw_paras:
        # Fallback: split by single newlines
        raw_paras = [p.strip() for p in plain_text.split("\n") if p.strip()]

    return [
        {"index": i + 1, "text": p, "chapter": "", "is_heading": False}
        for i, p in enumerate(raw_paras)
    ]
